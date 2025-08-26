import streamlit as st
import openai
import pdfplumber
from neo4j import GraphDatabase
import re
from docx import Document
import fitz  # PyMuPDF
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from fuzzywuzzy import fuzz  # Para calcular similitud de texto
from fuzzywuzzy import process # Para encontrar el mejor match
from collections import defaultdict

#Fase 2
# El objetivo de esta etpa es crear los titulos y subtitulos con su contenido, e inviarlos a Neo4j
#-------------------------------------------------------------------------------------------

# Configuración de OpenAI y Neo4j
openai.api_key = ""

# Configura la conexión a Neo4j
uri = "neo4j+s://.databases.neo4j.io"
username = "neo4j"
password = ""
driver = GraphDatabase.driver(uri, auth=(username, password))

#-------------------------------------------------------------------------------------------

# Función para extraer texto de PDF
def extract_text_from_pdf(pdf_file):
    text = ""
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

#-------------------------------------------------------------------------------------------

def extract_titles_and_subtitles(text):

    # Prompt para extraer títulos y subtítulos hasta el tercer nivel
    prompt = (
        "Extrae solo los títulos y subtítulos hasta el tercer nivel, del siguiente texto bibliografico,"
        "observa que el texto en la bibliografia utiliza un punto negro (•) antes del nombre del subtititulo del tercer nivel. "
        "Ordena los títulos en una estructura jerárquica con números (por ejemplo, "
        "1. Título Principal, 1.1 Subtítulo, 1.1.1 Subtitulo)  "
        "No incluyas información adicional sobre procedimientos o análisis. "
        "Texto:\n\n"
        f"{text}\n\n"
        "Títulos jerárquicos hasta el tercer nivel:"
    )

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Eres un asistente experto en procesamiento de texto bibliográfico."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=1000,
            temperature=0.3,
        )
        extracted_titles = response['choices'][0]['message']['content'].strip().split('\n')
        return extracted_titles
    except Exception as e:
        return [f"Error al procesar la solicitud: {e}"]

#-------------------------------------------------------------------------------------------

# Función para extraer descripciones de los títulos seleccionados
def extract_descriptions_for_selected_titles(text, selected_titles):
    extracted_descriptions = {}
    for title in selected_titles:
        prompt = (
            "Dado el siguiente texto y el titulo/subtitulo seleccionadoo, extrae su correspondiente descripción. El titulo/subtitulos seleccionado es:\n\n"
            f"{title}\n\n"
            "Y el Texto en donde se encuentran su correspondiente descripciones es:\n\n"
            f"{text}\n\n"
            "Extare exactamente la Descripción Completa segun el titulo/subtitulo seleccionado."
            "No comieces con una introduccion, replica exactamente lo solicitado, sin adiciones ni traducciones."
        )

        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Eres un asistente experto en procesamiento de documentos bibliográficos."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=2000,
                temperature=0.3,
            )
            description = response['choices'][0]['message']['content'].strip()
            extracted_descriptions[title] = description
        except Exception as e:
            # En caso de error, agregar el error como descripción para ese título
            extracted_descriptions[title] = f"Error al procesar la solicitud: {e}"
    return extracted_descriptions

#-------------------------------------------------------------------------------------------

# Función para crear nodos en Neo4j
def guardar_nodos_en_neo4j(titles_with_descriptions):
    """
    Guarda los títulos seleccionados como nodos en Neo4j con las propiedades 'nombre' y 'descripcion'.
    
    :param titles_with_descriptions: Diccionario con los títulos como claves y descripciones como valores.
    """
    nodos_con_ids = {}
    with driver.session() as session:
        for title, description in titles_with_descriptions.items():
            nodo_id = session.write_transaction(crear_nodo, title, description)
            nodos_con_ids[title] = nodo_id
    return nodos_con_ids

# Transacción para crear un nodo
def crear_nodo(tx, nombre, descripcion):
    query = (
        "MERGE (n:Título {nombre: $nombre}) "
        "SET n.descripcion = $descripcion "
        "RETURN ID(n) AS id"  
    )
    result = tx.run(query, nombre=nombre, descripcion=descripcion)
    record = result.single()  # Obtener el resultado único
    return record["id"] if record else None

# Cerrar conexión con Neo4j
def cerrar_conexion():
    driver.close()
#-------------------------------------------------------------------------------------------
#-------------------------------------------------------------------------------------------

# Fase 3
# El objetivo de esta etapa es realizar busquedas y establercer relaciones entre los nodos
#-------------------------------------------------------------------------------------------

def obtener_nodos_analisis_externos_flexibles(nombre_titulo, umbral=70):
    """
    Recupera los nodos de tipo 'Análisis' con el tipo 'externo' desde Neo4j,
    y realiza una comparación flexible con el nombre del nodo `Título` usando similitud textual.
    :param titulo_nombre: El nombre del nodo Título para comparar.
    :param umbral: Umbral de similitud (0-100) para considerar un nodo como similar.
    :return: Lista de diccionarios con ID, nombre y descripción de los nodos similares.
    """
    with driver.session() as session:
        # Obtener todos los nodos de análisis externos
        result = session.run(
            """
            MATCH (a:Análisis {tipo: 'externo'})
            RETURN ID(a) AS id, a.nombre AS nombre, a.descripcion AS descripcion
            """
        )
        # Lista de nodos con ID, nombre y descripción
        nodos_analisis = [
            {"id": record["id"], "nombre": record["nombre"], "descripcion": record["descripcion"]} 
            for record in result
        ]
        
    # Filtrar nodos por similitud flexible
    nodos_similares = []
    for nodo in nodos_analisis:
        similitud = fuzz.ratio(nombre_titulo.lower(), nodo["nombre"].lower())
        if similitud >= umbral:
            nodos_similares.append(
                {"id": nodo["id"], "nombre": nodo["nombre"], "descripcion": nodo["descripcion"], "similitud": similitud}
            )
    # Verificar si no se encontraron nodos similares
    if not nodos_similares:
        st.warning(f"No se encontraron nodos similares para el título: {nombre_titulo}")
    
    return nodos_similares    

#-------------------------------------------------------------------------------------------

def filtrar_nodos_por_palabras_comunes(descripcion_titulo, nodos_similares, umbral=0.1):
    """
    Filtra los nodos similares basándose en la cantidad de palabras comunes entre la descripción del título y la descripción del nodo.
    :param descripcion_titulo: Descripción del nodo Título.
    :param nodos_similares: Lista de nodos Análisis ya filtrados por similitud de nombres, con ID, nombre y descripción.
    :param umbral: Umbral mínimo de proporción de palabras comunes para considerar un nodo relevante.
    :return: Lista de nodos Análisis relevantes con ID, nombre y descripción.
    """
    if not nodos_similares:
        return []  # Si no hay nodos similares, devolver lista vacía
    
    nodos_relevantes = []
    palabras_titulo = set(descripcion_titulo.lower().split())  # Dividir en palabras únicas

    for nodo in nodos_similares:
        palabras_nodo = set(nodo["descripcion"].lower().split())  # Dividir descripción en palabras únicas
        palabras_comunes = palabras_titulo.intersection(palabras_nodo)  # Palabras comunes
        proporción = len(palabras_comunes) / len(palabras_titulo)  # Proporción de palabras comunes
        if proporción >= umbral:  # Si cumple con el umbral, añadir a nodos relevantes
            nodos_relevantes.append(nodo)
    
    return nodos_relevantes

#-------------------------------------------------------------------------------------------

def filtrar_nodos_con_verificacion(nodos_relevantes):
    """
    Filtra los nodos relevantes que tienen una relación VERIFICACION_DE_DESCRIPCION como nodos entrantes.
    :param nodos_relevantes: Lista de nodos relevantes con sus IDs, nombres y descripciones.
    :return: Lista de nodos con la relación VERIFICACION_DE_DESCRIPCION.
    """
    nodos_filtrados = []  # Lista para almacenar los nodos relacionados encontrados
    with driver.session() as session:
        for nodo in nodos_relevantes:
            # Consulta para verificar la relación VERIFICACION_DE_DESCRIPCION (entrante)
            result = session.run(
                """
                MATCH (n)-[:VERIFICACION_DE_DESCRIPCION]->(a)
                WHERE ID(a) = $id
                RETURN ID(n) AS id, n.nombre AS nombre, n.descripcion AS descripcion
                """,
                id=nodo["id"],
            )
            # Agregar los nodos relacionados encontrados a la lista
            nodos_filtrados.extend(
                [{"id": record["id"], "nombre": record["nombre"], "descripcion": record["descripcion"]} for record in result]
            )
    return nodos_filtrados

#-------------------------------------------------------------------------------------------

def crear_relacion_existe_semantica(titulo_id, nodo_id):
    """
    Crea la relación EXISTE_SEMANTICA entre un nodo Título y un nodo filtrado.
    :param titulo_id: ID del nodo Título.
    :param nodo_id: ID del nodo relacionado.
    """
    with driver.session() as session:
        session.run(
            """
            MATCH (t), (n)
            WHERE ID(t) = $titulo_id AND ID(n) = $nodo_id
            MERGE (n)-[:EXISTE_SEMANTICA]->(t)
            """,
            titulo_id=titulo_id,
            nodo_id=nodo_id,
        )

#-------------------------------------------------------------------------------------------
#-------------------------------------------------------------------------------------------

# Fase 3:
# Generacion y guardado
#-------------------------------------------------------------------------------------------

#Funcion para extraer los datos relacionados y generar una descripcion

def adaptar_descripciones_con_llm(titulo_id):
    """
    Adapta las descripciones del nodo 'Título' utilizando un LLM, basado en las descripciones relacionadas.

    Args:
        driver: Conexión activa a Neo4j.

    Returns:
        dict: Diccionario con títulos como claves y descripciones adaptadas como valores.
    """
    resultados_adaptados = {}

    with driver.session() as session:
        resultados = session.run(
            """
            MATCH (r)-[:EXISTE_SEMANTICA]->(t:`Título`)
            WHERE ID(t) = $titulo_id
            RETURN t.nombre AS titulo, t.descripcion AS titulo_descripcion, COLLECT(r.descripcion) AS descripciones_plantilla
            """,
            titulo_id=titulo_id,  
        )
        
        for registro in resultados:
            titulo = registro["titulo"]
            descripcion_original = registro["titulo_descripcion"]
            descripciones_plantilla = registro["descripciones_plantilla"]

            # Sin plantillas: traducir o copiar la descripción original
            if not descripciones_plantilla or len(descripciones_plantilla) == 0:
                prompt = (
                    f"Por favor, traduce la siguiente descripción al español o mantenla tal cual si ya está en español:\n\n"
                    f"{descripcion_original}"
                    f"No copies o pegues otras descripciones, solo traduce la descripción original."
                )
                try:
                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",  # Modelo definido directamente
                        messages=[
                            {"role": "system", "content": "Eres un asistente experto en traducción técnica."},
                            {"role": "user", "content": prompt},
                        ],
                        temperature=0.0,
                    )
                    descripcion_adaptada = response['choices'][0]['message']['content'].strip()
                    resultados_adaptados[titulo] = descripcion_adaptada

                except Exception as e:
                    st.warning(f"Error al procesar el título '{titulo}': {e}")
                    resultados_adaptados[titulo] = "Error en la adaptación."
                continue

            # Crear prompt para adaptaciones
            prompt = (
                f"Tarea:\n"
                f"Debes adaptar la siguiente descripción original a las plantillas proporcionadas. "
                f"No debes inventar nada nuevo. "
                f"Utiliza las plantillas como ejemplos para ajustar el estilo y contexto.\n\n"
                f"Descripción Original:\n{descripcion_original}\n\n"
                f"Plantillas:\n" + "\n".join(descripciones_plantilla) + "\n\n"
                f"Descripción Adaptada:"
                f"Si la diferencia entre la descripción original y las plantillas es muy grande, limitiate a traducir la descripción original EXACTAMENTE como es, no inventes nada."    
            )
            
            # Llamada al LLM
            try:
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",  # Modelo definido directamente
                    messages=[
                        {"role": "system", "content": "Eres un asistente experto en adaptar textos técnicos."},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.1,
                )
                descripcion_adaptada = response['choices'][0]['message']['content'].strip()
                resultados_adaptados[titulo] = descripcion_adaptada

            except Exception as e:
                st.warning(f"Error al procesar el título '{titulo}': {e}")
                resultados_adaptados[titulo] = "Error en la adaptación."

    return resultados_adaptados


#-------------------------------------------------------------------------------------------

#Funcion para guardar los datos en Neo4j



def guardar_descripciones_en_word(resultados_finales):
    """
    Guarda los títulos y descripciones en un archivo Word sin repetir los títulos en las descripciones.

    Args:
        resultados_finales (dict): Diccionario con títulos y descripciones adaptadas.

    Returns:
        str: Ruta del archivo Word generado.
    """
    doc = Document()
    doc.add_heading("Descripciones Adaptadas", level=1)

    # Iterar sobre los títulos principales en resultados_finales
    for titulo_principal, lista in resultados_finales.items():
        # Agregar el título principal al documento (nivel 2)
        doc.add_heading(titulo_principal, level=2)

        # Iterar sobre cada entrada en la lista de descripciones
        for entrada in lista:
            if "descripciones" in entrada:
                descripciones = entrada["descripciones"]
                # Agregar las descripciones sin repetir el título principal
                for subtitulo, descripcion in descripciones.items():
                    # Evitar repetir el título principal como subtítulo
                    if subtitulo.strip() != titulo_principal.strip():
                        doc.add_heading(subtitulo, level=3)  # Subtítulo
                    doc.add_paragraph(descripcion)  # Descripción adaptada

    # Guardar el archivo Word
    file_path = "descripciones_adaptadas.docx"
    doc.save(file_path)
    return file_path


#-------------------------------------------------------------------------------------------
#-------------------------------------------------------------------------------------------
#Interfaz con Streamlit
#-------------------------------------------------------------------------------------------
#-------------------------------------------------------------------------------------------
# Títulos y subtítulos extraídos
st.title("Extracción de Títulos y Subtítulos desde PDF")

# Entrada del usuario
uploaded_file = st.file_uploader("Sube un archivo PDF", type="pdf")

# Botón para procesar el archivo PDF
if uploaded_file:
    # Extraer texto del PDF
    text = extract_text_from_pdf(uploaded_file)

    # Guardar el texto completo y los títulos en session_state
    if "titles" not in st.session_state:
        titles = extract_titles_and_subtitles(text)
        st.session_state["titles"] = titles
        st.session_state["text"] = text  # Guardar el texto completo para usarlo luego en la extracción de descripciones

# Mostrar los títulos y checkboxes solo si ya se han extraído
if "titles" in st.session_state:
    st.write("Selecciona los títulos de interés:")

    # Inicializar el estado de los checkboxes si no existe
    if "checkbox_states" not in st.session_state:
        st.session_state["checkbox_states"] = {f"checkbox_{i}": False for i in range(len(st.session_state["titles"]))}

    # Lista para almacenar los títulos seleccionados
    selected_titles = []

    # Crear checkboxes persistentes para cada título
    for i, title in enumerate(st.session_state["titles"]):
        # Usar el estado del checkbox almacenado en session_state
        selected = st.checkbox(
            title, key=f"checkbox_{i}", value=st.session_state["checkbox_states"][f"checkbox_{i}"]
        )
        st.session_state["checkbox_states"][f"checkbox_{i}"] = selected

        # Agregar a la lista si está seleccionado
        if selected:
            selected_titles.append(title)

    # Guardar los títulos seleccionados en session_state
    st.session_state["selected_titles"] = selected_titles

    # Mostrar los títulos seleccionados
    st.write("Títulos seleccionados:")
    st.write(selected_titles)


#-------------------------------------------------------------------------------------------

# Descripcion extraida

# Mostrar títulos seleccionados y sus descripciones
if st.button("Mostrar Descripciones") and "titles" in st.session_state:
    st.write("Títulos seleccionados:", st.session_state.get("selected_titles", []))
    selected_titles = st.session_state["selected_titles"]

    # Validar si hay títulos seleccionados
    if selected_titles:
        # Llamar a la función para extraer las descripciones de los títulos seleccionados
        #st.write("Texto cargado en sesión:", st.session_state.get("text", "No disponible"))
        descriptions_dict = extract_descriptions_for_selected_titles(
            st.session_state["text"], selected_titles
        )

        # Guardar descriptions_dict en session_state
        st.session_state["descriptions_dict"] = descriptions_dict

        # Mostrar las descripciones en pantalla por título
        for title, description in descriptions_dict.items():
            st.subheader(f"Título: {title}")
            st.write(f"Descripción: {description}")
    else:
        st.warning("Por favor, selecciona al menos un título para ver su descripción.")

#-------------------------------------------------------------------------------------------

# Botón para crear relaciones y guardar en neo4j
if st.button("Crear Relaciones y guarda en neo4j"):
    if "selected_titles" in st.session_state and "descriptions_dict" in st.session_state:
        selected_titles = st.session_state["selected_titles"]
        descriptions_dict = st.session_state["descriptions_dict"]

        # Validar si hay títulos seleccionados
        if selected_titles:
            with st.spinner("Creando relaciones en Neo4j..."):
                try:
                    # Crear un diccionario de títulos y descripciones
                    titles_with_descriptions_dict = {
                        title: descriptions_dict[title] for title in selected_titles
                    }

                    # Guardar nodos en Neo4j y obtener sus IDs
                    nodos_con_ids = guardar_nodos_en_neo4j(titles_with_descriptions_dict)  # Retorna {title: neo4j_id}

                    # Crear una lista de títulos con nombres, descripciones e IDs de Neo4j
                    titles_with_descriptions = [
                        {
                            "nombre": title,
                            "descripcion": descriptions_dict[title],
                            "neo4j_id": nodos_con_ids[title],  # ID real de Neo4j
                        }
                        for title in selected_titles
                    ]


                    # Procesar cada título
                    resultados_titulos = {}  # Diccionario para almacenar los resultados de cada título
                    resultados_descripcion = {}
            
                    # Paso 1: Obtener nodos similares
                    for titulo in titles_with_descriptions:  # Iterar sobre los títulos seleccionados
                        nombre_titulo = titulo["nombre"]
                        descripcion_titulo = titulo["descripcion"]
                        nodos_similares = obtener_nodos_analisis_externos_flexibles(
                            nombre_titulo=nombre_titulo,
                            umbral=70  
                        )
                        # Almacenar nodos similares en el diccionario, usando el título como clave
                        resultados_titulos[nombre_titulo] = nodos_similares

                        # Paso 2: Filtrar por similitud en descripciones
                        try:
                            nodos_relevantes = filtrar_nodos_por_palabras_comunes(
                                descripcion_titulo= descripcion_titulo,
                                nodos_similares=nodos_similares,
                                umbral=0.3  # Umbral para palabras comunes
                            )
                            # Almacenar nodos relevantes
                            resultados_descripcion[nombre_titulo] = nodos_relevantes
                        except Exception as e:
                            st.error(f"Error al filtrar nodos para '{nombre_titulo}': {e}")
                            continue

                    # Paso 3: Obtener nodos relacionados con VERIFICACION_DE_DESCRIPCION
                    todos_los_nodos_con_verificacion = {}
                    for key, value_list in resultados_descripcion.items():  # Iterar sobre los títulos seleccionados
                        resultados_verificacion = []
                        for item in value_list:
                            nodos_con_verificacion = filtrar_nodos_con_verificacion([item])
                            resultados_verificacion.append(nodos_con_verificacion)
                        todos_los_nodos_con_verificacion[key] = resultados_verificacion

                    #Paso 4: Crear relaciones en Neo4j
                    for titles in titles_with_descriptions:
                        title_id = titles.get("neo4j_id")  # Obtener el ID de Neo4j del título
                        title_name = titles.get("nombre") 
                        if title_name in todos_los_nodos_con_verificacion:
                            nodos_relacionados = todos_los_nodos_con_verificacion[title_name]
                            for nodo_grupo in nodos_relacionados:  # Iterar sobre los grupos de nodos
                                for nodo in nodo_grupo:  # Iterar sobre los nodos dentro del grupo
                                    nodo_id = nodo.get("id")  # Obtener el ID de Neo4j del nodo
                                    if title_id and nodo_id:
                                        crear_relacion_existe_semantica(title_id, nodo_id)
                    st.write("Relaciones creadas con éxito.")
                                       
                except Exception as e:
                    st.error(f"Error al procesar las relaciones: {e}")
        else:
            st.warning("No hay títulos seleccionados para crear relaciones.")
    else:
        st.warning("No hay títulos seleccionados o descriptions_dict disponible en la sesión.")

#-------------------------------------------------------------------------------------------

#Fase 4: generacion y guardado
if "titles_with_descriptions" not in st.session_state:
    st.session_state["titles_with_descriptions"] = titles_with_descriptions
if st.button("Descargar archivo Word con descripciones adaptadas"):
    with st.spinner("Adaptando descripciones, esto puede tomar un momento..."):
        titles_with_descriptions = st.session_state["titles_with_descriptions"]
        resultados_finales = defaultdict(list)  # Cambiamos a defaultdict para evitar sobrescrituras
        
        # Iterar sobre títulos y obtener descripciones adaptadas
        for titles in titles_with_descriptions:
            title_id = titles.get("neo4j_id")
            resultados = adaptar_descripciones_con_llm(title_id)
            st.success(f"Adaptación completada para el título: {titles['nombre']}")
            
            # Agregar al diccionario usando append-like behavior
            resultados_finales[titles["nombre"]].append({
                "titulo": titles["nombre"],
                "descripciones": resultados
            })
        
        # Mostrar resultados finales en Streamlit (opcional)
        st.write(dict(resultados_finales))

    # Guardar y permitir descarga del archivo Word
    with st.spinner("Guardando descripciones adaptadas en un archivo Word..."):    
        archivo_word = guardar_descripciones_en_word(resultados_finales)
        if archivo_word is None:
            st.error("Error al guardar el archivo Word.")
        else:
            with open(archivo_word, "rb") as file:
                st.download_button(
                    label="Descargar Descripciones Adaptadas en Word",
                    data=file,
                    file_name=archivo_word,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
